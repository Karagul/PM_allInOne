from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import re
import pandas
from docx.shared import Inches

class create_docx_file():

    def __init__(self, fileName='New_Strikes_170718v2.docx'):

        self.fileName = fileName
        self.document = Document(self.fileName)
        self.tables = self.document.tables
        self.create_styles()
        self.mapping = {'rows' : self.tables[0].rows[2].cells[0].tables[0].rows,
                        'paragraphs' : self.tables[0].rows[2].cells[0].paragraphs,
                        'tables' : self.tables[0].rows[2].cells[0].tables}
        self.remove_old_table()
        # self.tables[0].rows[2].cells[0].width = Inches(13)

    def create_styles(self):
        try:
            styles = self.document.styles
            style = styles.add_style('table_style_black', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(10)
            style = styles.add_style('table_style_blue', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor(0, 0, 255)
            style = styles.add_style('table_style_bold', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(14)
            style.font.bold = True
            style = styles.add_style('table_style_normal', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(12)
        except ValueError:
            pass


    def remove_old_table(self):

        def move_table_after(table, paragraph):
            tbl, p = table._tbl, paragraph._p
            p.addnext(tbl)

        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)

        # print (self.tables[0].style)
        for table in self.mapping['tables']:
            for row in table.rows:
                remove_row(table, row)
            # table.add_column(5)
            cells = table.add_row().cells
            # # for cell in cells:
            #     # cell.width = Inches(10)
            # cells[0].merge(cells[3])
            # # cells[0].width = Inches(10)
            # table.style = 'Table Grid'
            # print (cells[0].width)
            for cell in cells:
                tempTable = cell.add_table(rows=2, cols=4)
                tempTable.style = 'Table Grid'
                tempTable.autofit = True
                break
            break
        # tempTable = self.tables[0].rows[2].cells[0].add_table(rows=2, cols=4)
        # tempTable.style = 'Table Grid'
        # tempTable.autofit = True
        # tempPara = self.tables[0].rows[2].cells[0].paragraphs[7]
        # move_table_after(tempTable, tempPara)

        # for table in self.mapping['tables']:
        #     # print (table)
        #     # for row in table.rows:
        #     #     remove_row(table, row)
        #     cells = table.add_row().cells

        # for table in self.tables[0].rows[2].cells[0].tables:
        #     print (table)
                


    def add_table_rows(self, rowData, textStyle):

        # for table in self.mapping['tables'][0].rows[0].cells[0].tables:

        #     cells = table.add_row().cells
        #     n = 0
        #     for cell in cells:
        #         paragraph = cell.paragraphs[0]
        #         paragraph.text = str(rowData[n])
        #         # print (rowData, paragraph.text)
        #         if textStyle == 'black':                
        #             paragraph.style = 'table_style_black'
        #             # print ('here black')
        #         elif textStyle == 'blue':
        #             paragraph.style = 'table_style_blue'
        #             # print ('here blue')
        #         n += 1
        #     table.style = 'Table Grid'
        #     table.autofit = True
        pass

    def adjust_number_date(self, tomorrow):

        today = (pandas.to_datetime('today').strftime('%Y-%m-%d'))

        text = self.mapping['paragraphs'][1].text
        enNumber = re.search(r'\d{3}(?=/17)', text).group()
        enNumber = str(int(enNumber) + 1)
        self.mapping['paragraphs'][1].text = re.sub(r'\d{3}(?=/17)', enNumber, text)
        self.mapping['paragraphs'][1].style = 'table_style_bold'

        text = self.mapping['paragraphs'][5].text
        # print (tomorrow)
        self.mapping['paragraphs'][5].text = re.sub(r'\d{4}-\d{2}-\d{2}', tomorrow, text)
        self.mapping['paragraphs'][5].style = 'table_style_normal'
        # print ('adjusted')


    def save_adjusted(self, fileName):

        self.document.save(fileName)

test = create_docx_file()
test.save_adjusted('New_Strikes_170718v2_FFFSSS.docx')