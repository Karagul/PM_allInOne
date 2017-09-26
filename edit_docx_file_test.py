from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import re
import pandas

class create_docx_file():

    def __init__(self, fileName):

        self.fileName = fileName
        self.document = Document(self.fileName)
        # self.tables = self.document.tables
        # self.create_styles()
        # self.tables[0].rows[2].cells[0].width = Inches(13)

    def create_styles(self):
        try:
            styles = self.document.styles
            style = styles.add_style('table_style_black', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(10)
            style = styles.add_style('table_style_blue', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor(0, 0, 255)
        except ValueError:
            pass

    def add_new_table(self, rowAmount):

        def move_table_after(table, paragraph):
            tbl, p = table._tbl, paragraph._p
            p.addnext(tbl)

        table = self.document.add_table(rows=rowAmount, cols=4)
        table.style = 'Table Grid'
        table.autofit = True

        paragraph = self.document.paragraphs[4]

        move_table_after(table, paragraph)

    def add_table_rows(self, rowNum, rowData, textStyle):

        table = self.document.tables[0]
        table_cells = table._cells
        for i in range(rowNum,(rowNum+1)):
            row_cells = table_cells[i*4:(i+1)*4]
            n = 0
            for cell in row_cells:
                paragraph = cell.paragraphs[0]
                paragraph.text = str(rowData[n])
                # print (rowData, paragraph.text)
                if textStyle == 'black':                
                    paragraph.style = 'table_style_black'
                    # print ('here black')
                elif textStyle == 'blue':
                    paragraph.style = 'table_style_blue'
                    # print ('here blue')
                n += 1
        table.autofit = True

    def adjust_number_date(self, enNumber, enYear, tomorrow):

        today = (pandas.to_datetime('today').strftime('%Y-%m-%d'))

        text = self.document.paragraphs[2].text

        finalId = enNumber + '/' + enYear
        self.document.paragraphs[2].text = re.sub('ENUMBER/YEAR', finalId, text)
        self.document.paragraphs[2].style = 'Para2'

        text = self.document.paragraphs[4].text
        # print (tomorrow)
        self.document.paragraphs[4].text = re.sub('LISTING_DATE', tomorrow, text)
        self.document.paragraphs[4].style = 'Para4'
        # print ('adjusted')


    def save_adjusted(self, fileName):

        self.document.save(fileName)